import mimetypes
import os
from wsgiref.util import FileWrapper

from django.http import StreamingHttpResponse
from django.shortcuts import render, redirect
from django.views.generic import DeleteView
from docx import Document
from docx.shared import Inches

from .models import Sellers, Buyers, Directory
from django.core.paginator import Paginator
from django.shortcuts import render

'''--------------------         CLASSES   (Sellers, Buyers)       --------------------'''


class DelBuyer(DeleteView):
    model = Buyers
    success_url = '/buyers'
    template_name = 'webApp/del_buyers.html'
    context_object_name = 'buyer'


class DelSeller(DeleteView):
    model = Sellers
    success_url = '/sellers'
    template_name = 'webApp/del_sellers.html'
    context_object_name = 'seller'


'''--------------------         Functions       --------------------'''


def head(request):
    num_apartments = Sellers.objects.all().count()
    num_buyers = Buyers.objects.all().count()

    districts = Directory.objects.all()
    rayon = {}
    for el in districts:
        rayon[el.name_district] = Sellers.objects.filter(id_district=el.id).count()
    num_district = len(districts)
    context = {
        'num_apartments': num_apartments,
        'num_buyers': num_buyers,
        'num_district': num_district,
        'area': rayon
    }
    return render(request, 'webApp/index.html', context)


'''--------------------         BEGIN SELLERS          --------------------'''


def get_sellers(request):
    sellers = Sellers.objects.order_by('-id')
    # подмена id района - названием
    for seller in sellers:
        seller.id_district = Directory.objects.filter(id=seller.id_district)[0].name_district
    paginator = Paginator(sellers, 6)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    return render(request, 'webApp/sellers.html', {'paginator': paginator, 'page_obj': page_obj, 'title': 'Продавцы'})


def add_apartment(request):
    error = ''
    if request.method == "POST":
        error = saveDataSeller(request, Sellers())
        if error == '':
            return redirect('sellers-page')

    return render(request, 'webApp/add_apartment.html', {"error": error})


def updateSellers(request, pk):
    sellers = Sellers.objects.get(id=pk)
    error = ''

    if request.method == "POST":
        error = saveDataSeller(request, sellers)
        if error == '':
            return redirect('sellers-page')

    form = {
        'fio': sellers.fio,
        'phone': sellers.phone,
        'district': Directory.objects.filter(id=sellers.id_district)[0].name_district,
        'number_floors': sellers.number_floors,
        'number_sell_floor': sellers.number_sell_floor,
        'area': sellers.area,
        'price': sellers.price
    }
    return render(request, 'webApp/add_apartment.html', {"form": form, "error": error})


def saveDataSeller(request, sellers):
    error = ''
    id_district = None
    try:
        sellers.fio = request.POST.get("fio")
        sellers.phone = request.POST.get("phone")
        sellers.number_floors = request.POST.get("number_floors")
        sellers.number_sell_floor = request.POST.get("number_sell_floor")
        sellers.area = request.POST.get("area")
        sellers.price = request.POST.get("price")

        nameDirectory = request.POST.get("district")
        resDistrict = Directory.objects.filter(name_district=nameDirectory)
        if len(resDistrict) > 0:
            sellers.id_district = resDistrict[0].id
        else:
            directory = Directory()
            directory.name_district = nameDirectory
            directory.save()
            id_district = Directory.objects.filter(name_district=nameDirectory)[0].id
            sellers.id_district = id_district
        sellers.save()

    except:
        Directory.objects.filter(id=id_district).delete()
        error = 'Форма была неверной!'

    return error


def sellersDetail(request, pk):
    sellers = Sellers.objects.get(id=pk)
    form = {
        'fio': sellers.fio,
        'phone': sellers.phone,
        'district': Directory.objects.filter(id=sellers.id_district)[0].name_district,
        'number_floors': sellers.number_floors,
        'number_sell_floor': sellers.number_sell_floor,
        'area': sellers.area,
        'price': sellers.price
    }
    return render(request, 'webApp/details_sellers.html', {"form": form})


'''--------------------         END SELLERS          --------------------'''
'''--------------------         BEGIN BUYERS          --------------------'''


def get_buyers(request):
    buyers = Buyers.objects.order_by('-id')
    # подмена id района - названием
    for buyer in buyers:
        buyer.id_district = Directory.objects.filter(id=buyer.id_district)[0].name_district
    paginator = Paginator(buyers, 4)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    return render(request, 'webApp/buyers.html', {'paginator': paginator, 'page_obj': page_obj, 'title': 'Покупатели'})


def add_buyer(request):
    error = ''
    if request.method == "POST":
        error = saveDataBuyers(request, Buyers())
        if error == '':
            return redirect('buyers-page')

    return render(request, 'webApp/add_buyer.html', {"error": error})


def saveDataBuyers(request, buyers):
    error = ''
    id_district = None
    try:
        buyers.fio = request.POST.get("fio")
        buyers.phone = request.POST.get("phone")
        buyers.min_area = request.POST.get("min_area")
        buyers.max_area = request.POST.get("max_area")
        buyers.price = request.POST.get("price")
        if request.POST.get("spicial_condition") == 'on':
            buyers.spicial_condition = True
        else:
            buyers.spicial_condition = False

        nameDirectory = request.POST.get("district")
        resDistrict = Directory.objects.filter(name_district=nameDirectory)
        if len(resDistrict) > 0:
            buyers.id_district = resDistrict[0].id
        else:
            directory = Directory()
            directory.name_district = nameDirectory
            directory.save()
            id_district = Directory.objects.filter(name_district=nameDirectory)[0].id
            buyers.id_district = id_district
        buyers.save()

    except:
        Directory.objects.filter(id=id_district).delete()
        error = 'Форма была неверной!'

    return error


def buyesDetail(request, pk):
    buyers = Buyers.objects.get(id=pk)
    form = {
        'fio': buyers.fio,
        'phone': buyers.phone,
        'district': Directory.objects.filter(id=buyers.id_district)[0].name_district,
        'min_area': buyers.min_area,
        'max_area': buyers.max_area,
        'price': buyers.price,
        'spicial_condition': buyers.spicial_condition
    }
    return render(request, 'webApp/details_buyers.html', {"form": form})


def updateBuyers(request, pk):
    buyers = Buyers.objects.get(id=pk)
    error = ''

    if request.method == "POST":
        error = saveDataBuyers(request, buyers)
        if error == '':
            return redirect('buyers-page')

    form = {
        'fio': buyers.fio,
        'phone': buyers.phone,
        'district': Directory.objects.filter(id=buyers.id_district)[0].name_district,
        'min_area': buyers.min_area,
        'max_area': buyers.max_area,
        'price': buyers.price,
        'spicial_condition': buyers.spicial_condition
    }
    return render(request, 'webApp/add_buyer.html', {"form": form, "error": error})


def createDocument(request, pk):
    buyers = Buyers.objects.get(id=pk)
    strSql = "SELECT * " \
             "FROM webApp_sellers as tblSell, webApp_buyers as tblBrs " \
             "WHERE tblSell.id_district = tblBrs.id_district " \
             "AND (tblSell.price * 1000000 + tblSell.price  * 1000000 * 0.03) <= tblBrs.price * 1000000 " \
             "AND tblSell.area BETWEEN tblBrs.min_area AND tblBrs.max_area;"
    sellersList = Sellers.objects.raw(strSql)

    listForm = []
    for seller in sellersList:
        if buyers.spicial_condition is False:
            if seller.number_sell_floor == 1 or seller.number_sell_floor == seller.number_floors:
                continue

        form_appartament = {
            'district': Directory.objects.filter(id=buyers.id_district)[0].name_district,
            'fio_sell': seller.fio,
            'phone_sell': seller.phone,
            'fio_buyer': buyers.fio,
            'phone_buyer': buyers.phone,
            'profit_agents': seller.price * 1000000 * 0.03
        }
        listForm.append(form_appartament)

    listForm = sortListForm(listForm)

    document = getDoc(listForm)
    document.save('Возможные варианты купли продажи.docx')

    response = StreamingHttpResponse(FileWrapper(open('Возможные варианты купли продажи.docx', 'rb'), 8192),
                                     content_type=mimetypes.guess_type('Возможные варианты купли продажи.docx')[0])
    response['Content-Length'] = os.path.getsize('Возможные варианты купли продажи.docx')
    response['Content-Disposition'] = "attachment; filename=%s" % 'Возможные варианты купли продажи.docx'
    return response

    #return redirect('buyers-page')


def getDoc(listForm):
    document = Document()

    for form in listForm:
        document.add_heading('Район ' + form['district'], 0)

        p = document.add_paragraph('ФИО продавца: ')
        p.add_run(form['fio_sell']).bold = True
        p = document.add_paragraph('Номер телефона: ')
        p.add_run(form['phone_sell']).bold = True
        p = document.add_paragraph('ФИО покупателя: ')
        p.add_run(form['fio_buyer']).bold = True
        p = document.add_paragraph('Номер телефона: ')
        p.add_run(form['phone_buyer']).bold = True
        p = document.add_paragraph('Прибыль агенства (3% от суммы сделки): ')
        p.add_run(str(form['profit_agents']) + "руб.").bold = True
    return document


def sortListForm(listForm):
    n = len(listForm)
    for i in range(n):
        for j in range(n-1):
            if listForm[j]['profit_agents'] < listForm[j + 1]['profit_agents']:
                temp = listForm[j]
                listForm[j] = listForm[j + 1]
                listForm[j + 1] = temp

    return listForm


'''--------------------         END BUYERS          --------------------'''
